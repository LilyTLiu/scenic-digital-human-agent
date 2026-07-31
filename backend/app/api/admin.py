"""
管理后台 API — 知识库 CRUD、对话记录、数据大屏
所有接口均对接 SQLite + ChromaDB 真实数据
"""
import json
import datetime
from typing import Optional
from fastapi import APIRouter, HTTPException, Depends, Query, Body
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import desc, func

from app.db.database import (
    get_db,
    KnowledgeDoc,
    ChatRecord,
    ScenicSpot,
    Feedback,
    VisitorReview,
    VisitorCheckin,
    LostFoundItem,
    ComplaintSuggestion,
    get_collection_name,
)
from app.core.rag import add_documents, delete_documents

router = APIRouter()


def _sync_chroma(doc: KnowledgeDoc, scenic_spot: str, db):
    """将一条知识文档切片后同步到 ChromaDB，返回新的 chroma_ids"""
    collection = get_collection_name(scenic_spot, db)

    # 删除旧片段
    old_ids = json.loads(doc.chroma_ids) if doc.chroma_ids else []
    if old_ids:
        delete_documents(collection, old_ids)

    # 添加新片段（ID 前缀确保唯一）
    added = add_documents(
        collection,
        [doc.content],
        [{"source": doc.title, "type": "kb_entry", "kb_id": doc.id}],
    )
    # ChromaDB 自动生成的 ID 格式为 chunk_0, chunk_1, ...
    # add_documents 返回的是 chunk 数量，不是 ID 列表
    # 需要从 collection 中获取刚添加的 ID（取最后 added 个）
    from app.core.rag import create_knowledge_base
    coll = create_knowledge_base(collection)
    all_ids = coll.get()["ids"]
    new_ids = all_ids[-added:] if added else []

    return new_ids


# ══════════════════════════════════════════════════════════════
# Pydantic Schemas
# ══════════════════════════════════════════════════════════════

class KnowledgeCreate(BaseModel):
    title: str
    content: str
    category: str = "通用"
    scenic_spot: str = "灵山胜境"


class KnowledgeUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    category: Optional[str] = None
    scenic_spot: Optional[str] = None


class KnowledgeOut(BaseModel):
    id: int
    title: str
    content: str
    category: str
    scenic_spot: str
    created_at: str
    updated_at: str


class PaginatedResponse(BaseModel):
    items: list[KnowledgeOut]
    total: int
    page: int
    size: int


class DashboardData(BaseModel):
    total_chats: int
    total_knowledge: int
    today_chats: int
    hot_keywords: list[str]
    daily_trend: list[dict] = []   # 最近7天每天对话量
    hourly_trend: list[dict] = []  # 今日每小时对话量
    total_sessions: int = 0
    spot_breakdown: list[dict] = []  # 各景点对话量+知识条目数


# ══════════════════════════════════════════════════════════════
# 知识库 CRUD
# ══════════════════════════════════════════════════════════════

@router.get("/knowledge", response_model=PaginatedResponse)
async def list_knowledge(
    scenic_spot: str = Query(default="灵山胜境"),
    category: str = Query(default=""),
    keyword: str = Query(default=""),
    page: int = Query(default=1, ge=1),
    size: int = Query(default=20, ge=1, le=100),
    db: Session = Depends(get_db),
):
    """知识库列表 — 支持分页、景区筛选、分类筛选、关键词搜索"""
    q = db.query(KnowledgeDoc)

    if scenic_spot:
        q = q.filter(KnowledgeDoc.scenic_spot == scenic_spot)
    if category:
        q = q.filter(KnowledgeDoc.category == category)
    if keyword:
        q = q.filter(
            (KnowledgeDoc.title.contains(keyword))
            | (KnowledgeDoc.content.contains(keyword))
        )

    total = q.count()
    rows = (
        q.order_by(desc(KnowledgeDoc.updated_at))
        .offset((page - 1) * size)
        .limit(size)
        .all()
    )

    return PaginatedResponse(
        items=[
            KnowledgeOut(
                id=r.id,
                title=r.title,
                content=r.content[:500],  # 列表只返回前 500 字
                category=r.category,
                scenic_spot=r.scenic_spot,
                created_at=r.created_at.isoformat() if r.created_at else "",
                updated_at=r.updated_at.isoformat() if r.updated_at else "",
            )
            for r in rows
        ],
        total=total,
        page=page,
        size=size,
    )


@router.get("/knowledge/{item_id}")
async def get_knowledge(item_id: int, db: Session = Depends(get_db)):
    """获取单条知识文档详情（含完整内容）"""
    doc = db.query(KnowledgeDoc).filter(KnowledgeDoc.id == item_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="知识条目不存在")
    return KnowledgeOut(
        id=doc.id,
        title=doc.title,
        content=doc.content,  # 详情返回完整内容
        category=doc.category,
        scenic_spot=doc.scenic_spot,
        created_at=doc.created_at.isoformat() if doc.created_at else "",
        updated_at=doc.updated_at.isoformat() if doc.updated_at else "",
    )


@router.post("/knowledge")
async def create_knowledge(data: KnowledgeCreate, db: Session = Depends(get_db)):
    """新增知识条目 → SQLite + ChromaDB 同步"""
    doc = KnowledgeDoc(
        title=data.title,
        content=data.content,
        category=data.category,
        scenic_spot=data.scenic_spot,
        created_at=datetime.datetime.now(),
        updated_at=datetime.datetime.now(),
    )
    db.add(doc)
    db.flush()  # 获取 doc.id

    # 同步到 ChromaDB
    try:
        new_ids = _sync_chroma(doc, data.scenic_spot, db)
        doc.chroma_ids = json.dumps(new_ids, ensure_ascii=False)
    except Exception:
        pass  # ChromaDB 同步失败不阻塞主流程

    db.commit()
    db.refresh(doc)

    return {
        "id": doc.id,
        "title": doc.title,
        "status": "created",
        "chroma_chunks": len(json.loads(doc.chroma_ids)) if doc.chroma_ids else 0,
    }


@router.put("/knowledge/{item_id}")
async def update_knowledge(
    item_id: int, data: KnowledgeUpdate, db: Session = Depends(get_db)
):
    """更新知识条目 → SQLite + ChromaDB 同步"""
    doc = db.query(KnowledgeDoc).filter(KnowledgeDoc.id == item_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="知识条目不存在")

    changed = False
    if data.title is not None:
        doc.title = data.title
        changed = True
    if data.content is not None:
        doc.content = data.content
        changed = True
    if data.category is not None:
        doc.category = data.category
        changed = True
    if data.scenic_spot is not None:
        doc.scenic_spot = data.scenic_spot
        changed = True

    if changed:
        doc.updated_at = datetime.datetime.now()
        # 内容变了才需要重建向量
        if data.content is not None:
            try:
                new_ids = _sync_chroma(doc, doc.scenic_spot, db)
                doc.chroma_ids = json.dumps(new_ids, ensure_ascii=False)
            except Exception:
                pass

    db.commit()
    db.refresh(doc)

    return {
        "id": doc.id,
        "title": doc.title,
        "status": "updated",
        "chroma_chunks": len(json.loads(doc.chroma_ids)) if doc.chroma_ids else 0,
    }


@router.delete("/knowledge/{item_id}")
async def delete_knowledge(item_id: int, db: Session = Depends(get_db)):
    """删除知识条目 → ChromaDB + SQLite"""
    doc = db.query(KnowledgeDoc).filter(KnowledgeDoc.id == item_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="知识条目不存在")

    # 从 ChromaDB 删除
    try:
        collection = get_collection_name(doc.scenic_spot, db)
        old_ids = json.loads(doc.chroma_ids) if doc.chroma_ids else []
        if old_ids:
            delete_documents(collection, old_ids)
    except Exception:
        pass

    db.delete(doc)
    db.commit()

    return {"deleted": item_id, "status": "deleted"}


# ══════════════════════════════════════════════════════════════
# 知识库分类列表（前端下拉菜单用）
# ══════════════════════════════════════════════════════════════

@router.get("/knowledge-categories")
async def list_categories(scenic_spot: str = Query(default="灵山胜境"), db: Session = Depends(get_db)):
    """返回当前景区下所有分类"""
    cats = (
        db.query(KnowledgeDoc.category)
        .filter(KnowledgeDoc.scenic_spot == scenic_spot)
        .distinct()
        .all()
    )
    return {"categories": [c[0] for c in cats if c[0]]}


# ══════════════════════════════════════════════════════════════
# 对话记录
# ══════════════════════════════════════════════════════════════

@router.get("/chat-records")
async def list_chat_records(
    scenic_spot: str = Query(default=""),
    session_id: str = Query(default=""),
    page: int = Query(default=1, ge=1),
    size: int = Query(default=50, ge=1, le=200),
    db: Session = Depends(get_db),
):
    """对话记录列表 — 支持按景区/会话筛选"""
    q = db.query(ChatRecord)
    if scenic_spot:
        q = q.filter(ChatRecord.scenic_spot == scenic_spot)
    if session_id:
        q = q.filter(ChatRecord.session_id == session_id)

    total = q.count()
    rows = (
        q.order_by(desc(ChatRecord.created_at))
        .offset((page - 1) * size)
        .limit(size)
        .all()
    )

    return {
        "items": [
            {
                "id": r.id,
                "session_id": r.session_id,
                "scenic_spot": r.scenic_spot,
                "user_input": r.user_input,
                "ai_reply": r.ai_reply[:300] if r.ai_reply else "",
                "created_at": r.created_at.isoformat() if r.created_at else "",
            }
            for r in rows
        ],
        "total": total,
        "page": page,
        "size": size,
    }


@router.get("/chat-sessions")
async def list_sessions(db: Session = Depends(get_db)):
    """返回不重复的 session_id 列表（供下拉筛选）"""
    sessions = db.query(ChatRecord.session_id).distinct().all()
    return {"sessions": [s[0] for s in sessions if s[0]]}


# ══════════════════════════════════════════════════════════════
# 数据大屏
# ══════════════════════════════════════════════════════════════

@router.get("/dashboard", response_model=DashboardData)
async def get_dashboard(db: Session = Depends(get_db)):
    """数据大屏 — 真实统计数据 + 趋势图表"""
    total_chats = db.query(func.count(ChatRecord.id)).scalar() or 0
    total_knowledge = db.query(func.count(KnowledgeDoc.id)).scalar() or 0
    total_sessions = db.query(func.count(func.distinct(ChatRecord.session_id))).scalar() or 0

    today = datetime.date.today()
    today_str = today.isoformat()
    today_chats = (
        db.query(func.count(ChatRecord.id))
        .filter(func.date(ChatRecord.created_at) == today_str)
        .scalar()
        or 0
    )

    # 最近 7 天趋势
    daily_trend = []
    for i in range(6, -1, -1):
        d = today - datetime.timedelta(days=i)
        d_str = d.isoformat()
        count = (
            db.query(func.count(ChatRecord.id))
            .filter(func.date(ChatRecord.created_at) == d_str)
            .scalar()
            or 0
        )
        daily_trend.append({"date": d.strftime("%m-%d"), "count": count})

    # 今日每小时分布
    hourly_trend = []
    for h in range(8, 22):  # 8:00-21:00
        count = (
            db.query(func.count(ChatRecord.id))
            .filter(func.date(ChatRecord.created_at) == today_str)
            .filter(func.strftime("%H", ChatRecord.created_at) == f"{h:02d}")
            .scalar()
            or 0
        )
        hourly_trend.append({"hour": f"{h}:00", "count": count})

    # 热门关键词：按问题频率
    hot_keywords: list[str] = []
    rows = (
        db.query(ChatRecord.user_input, func.count(ChatRecord.id).label("cnt"))
        .group_by(ChatRecord.user_input)
        .order_by(desc("cnt"))
        .limit(10)
        .all()
    )
    hot_keywords = [
        (r[0][:20] + ("..." if len(r[0]) > 20 else "")) if r[0] else ""
        for r in rows
    ]

    # 各景点对话量 + 知识覆盖
    spot_chat_rows = (
        db.query(ChatRecord.scenic_spot, func.count(ChatRecord.id).label("cnt"))
        .filter(ChatRecord.scenic_spot != "", ChatRecord.scenic_spot.isnot(None))
        .group_by(ChatRecord.scenic_spot)
        .order_by(desc("cnt"))
        .all()
    )
    spot_kb_rows = dict(
        db.query(KnowledgeDoc.scenic_spot, func.count(KnowledgeDoc.id))
        .filter(KnowledgeDoc.scenic_spot != "", KnowledgeDoc.scenic_spot.isnot(None))
        .group_by(KnowledgeDoc.scenic_spot)
        .all()
    )
    spot_breakdown = [
        {
            "spot_id": r.scenic_spot, "chat_count": r.cnt,
            "kb_count": spot_kb_rows.get(r.scenic_spot, 0),
        }
        for r in spot_chat_rows
    ]

    return DashboardData(
        total_chats=total_chats,
        total_knowledge=total_knowledge,
        today_chats=today_chats,
        hot_keywords=hot_keywords,
        daily_trend=daily_trend,
        hourly_trend=hourly_trend,
        total_sessions=total_sessions,
        spot_breakdown=spot_breakdown,
    )


# ══════════════════════════════════════════════════════════════
# 数字人管理（保留 Mock，后续开发）
# ══════════════════════════════════════════════════════════════

@router.get("/digital-humans")
async def list_digital_humans():
    return {
        "active": "default",
        "humans": [
            {"id": "default", "name": "导游小文", "voice": "zh-CN-XiaoxiaoNeural", "provider": "xmov"},
            {"id": "guide-yun", "name": "导游小云", "voice": "zh-CN-XiaoxiaoNeural", "provider": "xmov"},
            {"id": "guide-ling", "name": "导游小灵", "voice": "zh-CN-XiaoxiaoNeural", "provider": "xmov"},
        ],
    }


@router.put("/digital-humans/{human_id}")
async def update_digital_human(human_id: str, config: dict):
    if human_id not in {"default", "guide-yun", "guide-ling"}:
        raise HTTPException(status_code=404, detail="数字人不存在")
    return {"success": True, "active": human_id, "id": human_id, **config}


class SwitchPersonaRequest(BaseModel):
    persona: str = "miaoyin"


@router.post("/switch-persona")
async def switch_persona(req: SwitchPersonaRequest):
    return {"success": True, "persona": req.persona}


# ══════════════════════════════════════════════════════════════
# 景区管理
# ══════════════════════════════════════════════════════════════

class ScenicSpotCreate(BaseModel):
    name: str
    slug: str
    description: str = ""


class ScenicSpotOut(BaseModel):
    id: int
    name: str
    slug: str
    description: str
    enabled: int


@router.get("/scenic-spots")
async def list_scenic_spots(
    include_disabled: bool = False,
    db: Session = Depends(get_db),
):
    """景区列表 — 前端下拉菜单和 Admin 管理用"""
    q = db.query(ScenicSpot)
    if not include_disabled:
        q = q.filter(ScenicSpot.enabled == 1)
    spots = q.order_by(ScenicSpot.id).all()
    return {
        "items": [
            ScenicSpotOut(
                id=s.id,
                name=s.name,
                slug=s.slug,
                description=s.description or "",
                enabled=s.enabled or 1,
            )
            for s in spots
        ]
    }


@router.post("/scenic-spots")
async def create_scenic_spot(data: ScenicSpotCreate, db: Session = Depends(get_db)):
    """新增景区 — 自动创建对应 ChromaDB 集合"""
    existing = db.query(ScenicSpot).filter(
        (ScenicSpot.slug == data.slug) | (ScenicSpot.name == data.name)
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail=f"景区 '{data.slug}' 或名称已存在")

    spot = ScenicSpot(
        name=data.name,
        slug=data.slug,
        description=data.description,
        enabled=1,
        created_at=datetime.datetime.now(),
    )
    db.add(spot)

    # 预创建 ChromaDB 集合
    try:
        from app.core.rag import create_knowledge_base
        create_knowledge_base(data.slug)
    except Exception:
        pass

    db.commit()
    db.refresh(spot)
    return ScenicSpotOut(id=spot.id, name=spot.name, slug=spot.slug, description=spot.description or "", enabled=1)


@router.put("/scenic-spots/{spot_id}")
async def update_scenic_spot(spot_id: int, data: ScenicSpotCreate, db: Session = Depends(get_db)):
    """更新景区信息"""
    spot = db.query(ScenicSpot).filter(ScenicSpot.id == spot_id).first()
    if not spot:
        raise HTTPException(status_code=404, detail="景区不存在")
    spot.name = data.name
    spot.slug = data.slug
    spot.description = data.description
    db.commit()
    db.refresh(spot)
    return ScenicSpotOut(id=spot.id, name=spot.name, slug=spot.slug, description=spot.description or "", enabled=spot.enabled or 1)


@router.delete("/scenic-spots/{spot_id}")
async def delete_scenic_spot(spot_id: int, db: Session = Depends(get_db)):
    """删除景区（软删除 = 设为停用）"""
    spot = db.query(ScenicSpot).filter(ScenicSpot.id == spot_id).first()
    if not spot:
        raise HTTPException(status_code=404, detail="景区不存在")
    spot.enabled = 0
    db.commit()
    return {"deleted": spot_id, "status": "disabled"}


# ══════════════════════════════════════════════════════════════
# 游客反馈
# ══════════════════════════════════════════════════════════════

class FeedbackRequest(BaseModel):
    rating: int  # 1=点赞, -1=踩
    question: str = ""  # 触发这条回复的用户问题


# ══════════════════════════════════════════════════════════════
# 游客行为分析 & 推荐（基于 140K 数据集）
# ══════════════════════════════════════════════════════════════

import sqlite3 as _sqlite3
import os as _os

_TOURIST_DB = _os.path.join(_os.path.dirname(__file__), "..", "..", "tourist_behavior.db")


def _query_tourist_db(query: str, params=()):
    """查询游客行为数据库"""
    if not _os.path.exists(_TOURIST_DB):
        return []
    conn = _sqlite3.connect(_TOURIST_DB)
    conn.row_factory = _sqlite3.Row
    try:
        return [dict(r) for r in conn.execute(query, params).fetchall()]
    finally:
        conn.close()


@router.get("/tourist/insights")
async def tourist_insights():
    """游客行为数据洞察"""
    rows = _query_tourist_db("""
        SELECT attraction_name, COUNT(*) as visits,
               ROUND(AVG(satisfaction), 1) as avg_sat,
               ROUND(AVG(total_cost), 0) as avg_cost,
               ROUND(AVG(age), 0) as avg_age
        FROM tourist_behaviors
        WHERE attraction_name LIKE '%灵山%' OR attraction_name LIKE '%无锡%' OR attraction_name LIKE '%太湖%'
        GROUP BY attraction_name
        ORDER BY visits DESC LIMIT 10
    """)
    return {"insights": rows}


@router.get("/tourist/recommend")
async def recommend_attractions(
    age: int = 30,
    budget: str = "medium",  # low / medium / high
    style: str = "",  # 深度游/亲子游/轻松游
    group: str = "couple",  # solo/couple/family/friends
):
    """基于用户画像推荐景点 + 预估花费"""
    # 预算范围映射
    budget_range = {"low": (0, 200), "medium": (200, 600), "high": (600, 9999)}
    lo, hi = budget_range.get(budget, (200, 600))

    # 按满意度 + 访问量排序，筛选相近年龄段 + 预算范围
    rows = _query_tourist_db("""
        SELECT attraction_name, attraction_type,
               COUNT(*) as visits,
               ROUND(AVG(satisfaction), 1) as avg_sat,
               ROUND(AVG(total_cost), 0) as est_cost,
               ROUND(AVG(stay_duration), 1) as avg_hours
        FROM tourist_behaviors
        WHERE age BETWEEN ? AND ?
          AND total_cost BETWEEN ? AND ?
          AND (attraction_type NOT IN ('', 'N/A'))
        GROUP BY attraction_name
        ORDER BY avg_sat DESC, visits DESC
        LIMIT 20
    """, (max(10, age - 10), min(70, age + 10), lo, hi))

    # 分组类型适配
    group_filter = {"solo": 1, "couple": 2, "family": ">=3", "friends": ">=2"}
    gf = group_filter.get(group, 2)

    return {
        "recommendations": [
            {"name": r["attraction_name"], "type": r["attraction_type"],
             "visits": r["visits"], "satisfaction": r["avg_sat"],
             "est_cost": r["est_cost"], "avg_hours": r["avg_hours"]}
            for r in rows[:12]
        ]
    }


# ══════════════════════════════════════════════════════════════
# 一键导入示例资料
# ══════════════════════════════════════════════════════════════

DEMO_DATA_DIR = _os.path.join(_os.path.dirname(__file__), "..", "..", "..", "示范景区公开资料包")


@router.post("/import-demo")
async def import_demo_data(db: Session = Depends(get_db)):
    """解析示范景区公开资料包中的docx文件，导入知识库+向量化"""
    try:
        from docx import Document
    except ImportError:
        return {"status": "error", "error": "后端缺少python-docx库"}

    imported = 0
    collection = get_collection_name("灵山胜境", db)

    # 1. 导入结构化景点数据集
    struct_file = _os.path.join(DEMO_DATA_DIR, "灵山胜境 景点结构化数据集.docx")
    if _os.path.exists(struct_file):
        try:
            doc = Document(struct_file)
            for table in doc.tables:
                if not table.rows: continue
                header = [c.text.strip() for c in table.rows[0].cells]
                col_map = {}
                for idx, h in enumerate(header):
                    for key in ["景区名称","景点名称","景点位置","外观","核心功能","文化内涵","详细介绍","最佳打卡点","开放","备注"]:
                        if key in h and key not in col_map: col_map[key] = idx
                name_idx = col_map.get("景点名称")
                if name_idx is None: continue
                scenic_idx = col_map.get("景区名称", name_idx)
                for row in table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) <= name_idx or not cells[name_idx]: continue
                    name = cells[name_idx]
                    scenic = cells[scenic_idx] if scenic_idx < len(cells) else "灵山胜境"
                    parts = []
                    for key, ci in col_map.items():
                        if ci < len(cells) and cells[ci] and key not in ("景区名称","景点名称"):
                            parts.append(f"【{key}】{cells[ci]}")
                    content = "\n".join(parts) or name
                    existing = db.query(KnowledgeDoc).filter(KnowledgeDoc.title == name).first()
                    if not existing:
                        db.add(KnowledgeDoc(title=name, content=content[:5000], category="景点讲解",
                            scenic_spot=scenic, updated_at=datetime.datetime.now(), chroma_ids="[]"))
                        db.commit()
                    try:
                        cnt = add_documents(collection, [content], [{"source": name, "category": "景点讲解"}])
                        if existing:
                            from app.core.rag import create_knowledge_base as _cb
                            all_ids = _cb(collection).get()["ids"]
                            existing.chroma_ids = json.dumps(all_ids[-cnt:] if cnt else [])
                    except Exception: pass
                    imported += 1
        except Exception as e:
            return {"status": "error", "error": f"解析结构化数据集失败: {e}"}

    # 2. 导入综合指南
    guide_file = _os.path.join(DEMO_DATA_DIR, "灵山胜境：历史、文化、景点特色与个性化游览指南.docx")
    if _os.path.exists(guide_file):
        try:
            doc = Document(guide_file)
            current_title, current_text = "", ""
            def flush():
                nonlocal current_title, current_text, imported
                if current_title and current_text.strip():
                    existing = db.query(KnowledgeDoc).filter(KnowledgeDoc.title == current_title).first()
                    if not existing:
                        db.add(KnowledgeDoc(title=current_title, content=current_text.strip()[:5000],
                            category="文史资料", scenic_spot="灵山胜境", updated_at=datetime.datetime.now(), chroma_ids="[]"))
                        db.commit()
                    try: add_documents(collection, [current_text.strip()], [{"source": current_title, "category": "文史资料"}])
                    except Exception: pass
                    imported += 1
                current_title = ""; current_text = ""
            for p in doc.paragraphs:
                txt = p.text.strip()
                if not txt: continue
                is_heading = bool(p.style and p.style.name and "Heading" in p.style.name)
                if is_heading or (len(txt) <= 30 and not txt.endswith(("。","！","？"))):
                    flush()
                    current_title = txt
                else:
                    current_text += txt + "\n"
            flush()
        except Exception as e:
            return {"status": "error", "error": f"解析综合指南失败: {e}"}

    return {"status": "ok", "imported": imported, "message": f"成功导入 {imported} 条知识"}


@router.post("/feedback")
async def submit_feedback(data: FeedbackRequest, db: Session = Depends(get_db)):
    """游客点赞/踩"""
    fb = Feedback(rating=data.rating, question=data.question, created_at=datetime.datetime.now())
    db.add(fb)
    db.commit()
    return {"status": "ok", "rating": data.rating}


@router.get("/feedback/stats")
async def feedback_stats(db: Session = Depends(get_db)):
    """反馈统计"""
    total = db.query(func.count(Feedback.id)).scalar() or 0
    likes = db.query(func.count(Feedback.id)).filter(Feedback.rating == 1).scalar() or 0
    dislikes = db.query(func.count(Feedback.id)).filter(Feedback.rating == -1).scalar() or 0
    recent = db.query(Feedback).order_by(desc(Feedback.created_at)).limit(10).all()
    suggestion_total = db.query(func.count(ComplaintSuggestion.id)).filter(ComplaintSuggestion.deleted == 0).scalar() or 0
    suggestion_pending = db.query(func.count(ComplaintSuggestion.id)).filter(
        ComplaintSuggestion.deleted == 0,
        ComplaintSuggestion.status == "pending",
    ).scalar() or 0
    suggestion_recent = (
        db.query(ComplaintSuggestion)
        .filter(ComplaintSuggestion.deleted == 0)
        .order_by(desc(ComplaintSuggestion.created_at))
        .limit(10)
        .all()
    )
    return {
        "total": total, "likes": likes, "dislikes": dislikes,
        "rate": round(likes / total * 100, 1) if total > 0 else 0,
        "recent": [{"id": f.id, "rating": f.rating, "question": f.question or "",
                     "created_at": f.created_at.isoformat() if f.created_at else ""} for f in recent],
        "suggestions_total": suggestion_total,
        "suggestions_pending": suggestion_pending,
        "suggestions_recent": [
            {
                "id": s.id,
                "category": s.category,
                "name": s.name,
                "contact": s.contact,
                "content": s.content,
                "status": s.status,
                "created_at": s.created_at.isoformat() if s.created_at else "",
            }
            for s in suggestion_recent
        ],
    }


# ══════════════════════════════════════════════════════════════
# 失物招领 / 投诉建议
# ══════════════════════════════════════════════════════════════

class LostFoundCreate(BaseModel):
    item_name: str
    item_type: str = "lost"
    location: str = ""
    description: str = ""
    contact: str = ""


class LostFoundUpdate(BaseModel):
    item_name: Optional[str] = None
    item_type: Optional[str] = None
    location: Optional[str] = None
    description: Optional[str] = None
    contact: Optional[str] = None
    status: Optional[str] = None


class SuggestionCreate(BaseModel):
    category: str = "suggestion"
    name: str = ""
    contact: str = ""
    content: str


class SuggestionUpdate(BaseModel):
    status: Optional[str] = None


@router.get("/lost-found")
async def list_lost_found(
    page: int = Query(default=1, ge=1),
    size: int = Query(default=50, ge=1, le=200),
    status: str = Query(default=""),
    db: Session = Depends(get_db),
):
    q = db.query(LostFoundItem).filter(LostFoundItem.deleted == 0)
    if status:
        q = q.filter(LostFoundItem.status == status)
    total = q.count()
    rows = q.order_by(desc(LostFoundItem.created_at)).offset((page - 1) * size).limit(size).all()
    return {
        "items": [
            {
                "id": item.id,
                "item_name": item.item_name,
                "item_type": item.item_type,
                "location": item.location,
                "description": item.description,
                "contact": item.contact,
                "status": item.status,
                "created_at": item.created_at.isoformat() if item.created_at else "",
            }
            for item in rows
        ],
        "total": total,
        "page": page,
        "size": size,
    }


@router.post("/lost-found")
async def create_lost_found(data: LostFoundCreate, db: Session = Depends(get_db)):
    if not data.item_name.strip():
        raise HTTPException(status_code=400, detail="物品名称不能为空")
    item_type = data.item_type if data.item_type in {"lost", "found"} else "lost"
    item = LostFoundItem(
        item_name=data.item_name.strip(),
        item_type=item_type,
        location=data.location.strip(),
        description=data.description.strip(),
        contact=data.contact.strip(),
        status="open",
        created_at=datetime.datetime.now(),
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"status": "ok", "id": item.id}


@router.put("/lost-found/{item_id}")
async def update_lost_found(item_id: int, data: LostFoundUpdate, db: Session = Depends(get_db)):
    item = db.query(LostFoundItem).filter(LostFoundItem.id == item_id, LostFoundItem.deleted == 0).first()
    if not item:
        raise HTTPException(status_code=404, detail="失物招领条目不存在")
    for field in ["item_name", "item_type", "location", "description", "contact", "status"]:
        value = getattr(data, field)
        if value is not None:
            setattr(item, field, value.strip() if isinstance(value, str) else value)
    if item.item_type not in {"lost", "found"}:
        item.item_type = "lost"
    if item.status not in {"open", "claimed", "closed"}:
        item.status = "open"
    db.commit()
    return {"status": "ok", "id": item.id}


@router.delete("/lost-found/{item_id}")
async def delete_lost_found(item_id: int, db: Session = Depends(get_db)):
    item = db.query(LostFoundItem).filter(LostFoundItem.id == item_id, LostFoundItem.deleted == 0).first()
    if not item:
        raise HTTPException(status_code=404, detail="失物招领条目不存在")
    item.deleted = 1
    db.commit()
    return {"deleted": item_id, "status": "soft-deleted"}


@router.get("/suggestions")
async def list_suggestions(
    page: int = Query(default=1, ge=1),
    size: int = Query(default=50, ge=1, le=200),
    status: str = Query(default=""),
    db: Session = Depends(get_db),
):
    q = db.query(ComplaintSuggestion).filter(ComplaintSuggestion.deleted == 0)
    if status:
        q = q.filter(ComplaintSuggestion.status == status)
    total = q.count()
    rows = q.order_by(desc(ComplaintSuggestion.created_at)).offset((page - 1) * size).limit(size).all()
    return {
        "items": [
            {
                "id": item.id,
                "category": item.category,
                "name": item.name,
                "contact": item.contact,
                "content": item.content,
                "status": item.status,
                "created_at": item.created_at.isoformat() if item.created_at else "",
            }
            for item in rows
        ],
        "total": total,
        "page": page,
        "size": size,
    }


@router.post("/suggestions")
async def create_suggestion(data: SuggestionCreate, db: Session = Depends(get_db)):
    if not data.content.strip():
        raise HTTPException(status_code=400, detail="意见内容不能为空")
    category = data.category if data.category in {"suggestion", "complaint"} else "suggestion"
    item = ComplaintSuggestion(
        category=category,
        name=data.name.strip(),
        contact=data.contact.strip(),
        content=data.content.strip(),
        status="pending",
        created_at=datetime.datetime.now(),
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"status": "ok", "id": item.id}


@router.put("/suggestions/{item_id}")
async def update_suggestion(item_id: int, data: SuggestionUpdate, db: Session = Depends(get_db)):
    item = db.query(ComplaintSuggestion).filter(
        ComplaintSuggestion.id == item_id,
        ComplaintSuggestion.deleted == 0,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="意见不存在")
    if data.status is not None:
        item.status = data.status if data.status in {"pending", "processing", "resolved"} else "pending"
    db.commit()
    return {"status": "ok", "id": item.id}


@router.delete("/suggestions/{item_id}")
async def delete_suggestion(item_id: int, db: Session = Depends(get_db)):
    item = db.query(ComplaintSuggestion).filter(
        ComplaintSuggestion.id == item_id,
        ComplaintSuggestion.deleted == 0,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="意见不存在")
    item.deleted = 1
    db.commit()
    return {"deleted": item_id, "status": "soft-deleted"}


# ══════════════════════════════════════════════════════════════
# 游客评价 (Visitor Reviews)
# ══════════════════════════════════════════════════════════════

class ReviewCreate(BaseModel):
    spot_id: str
    author: str = "灵山游客"
    avatar: str = "😊"
    rating: int = 5  # 1-5
    text: str


class ReviewOut(BaseModel):
    id: int
    spot_id: str
    author: str
    avatar: str
    rating: int
    text: str
    time: str  # 日期字符串 YYYY-MM-DD


class CheckinCreate(BaseModel):
    spot_id: str
    author: str = "灵山游客"
    image: str = ""
    caption: str


class CheckinOut(BaseModel):
    id: int
    spot_id: str
    author: str
    image: str
    caption: str
    time: str


@router.get("/reviews")
async def list_reviews(
    spot_id: str = Query(default=""),
    page: int = Query(default=1, ge=1),
    size: int = Query(default=50, ge=1, le=200),
    db: Session = Depends(get_db),
):
    """获取游客评价列表 — 可按景点筛选"""
    q = db.query(VisitorReview).filter(VisitorReview.deleted == 0)
    if spot_id:
        q = q.filter(VisitorReview.spot_id == spot_id)
    total = q.count()
    rows = q.order_by(desc(VisitorReview.created_at)).offset((page - 1) * size).limit(size).all()
    return {
        "items": [
            ReviewOut(
                id=r.id, spot_id=r.spot_id, author=r.author, avatar=r.avatar,
                rating=r.rating, text=r.text,
                time=r.created_at.strftime("%Y-%m-%d") if r.created_at else "",
            ).dict() | {
                "created_at": r.created_at.isoformat() if r.created_at else "",
            } for r in rows
        ],
        "total": total, "page": page, "size": size,
    }


@router.post("/reviews")
async def create_review(data: ReviewCreate, db: Session = Depends(get_db)):
    """提交新评价"""
    if not data.text.strip():
        raise HTTPException(status_code=400, detail="评价内容不能为空")
    if not data.author.strip():
        data.author = "灵山游客"
    rev = VisitorReview(
        spot_id=data.spot_id, author=data.author.strip(), avatar=data.avatar,
        rating=max(1, min(5, data.rating)), text=data.text.strip(),
        created_at=datetime.datetime.now(),
    )
    db.add(rev)
    db.commit()
    db.refresh(rev)
    return {
        "id": rev.id, "spot_id": rev.spot_id, "author": rev.author,
        "avatar": rev.avatar, "rating": rev.rating, "text": rev.text,
        "time": rev.created_at.strftime("%Y-%m-%d") if rev.created_at else "",
    }


@router.delete("/reviews/{review_id}")
async def delete_review(review_id: int, db: Session = Depends(get_db)):
    """删除评价（管理后台用）"""
    rev = db.query(VisitorReview).filter(VisitorReview.id == review_id, VisitorReview.deleted == 0).first()
    if not rev:
        raise HTTPException(status_code=404, detail="评价不存在")
    rev.deleted = 1
    db.commit()
    return {"deleted": review_id, "status": "soft-deleted"}


# ══════════════════════════════════════════════════════════════
# 游客打卡 (Visitor Checkins)
# ══════════════════════════════════════════════════════════════

@router.get("/checkins")
async def list_checkins(
    spot_id: str = Query(default=""),
    page: int = Query(default=1, ge=1),
    size: int = Query(default=50, ge=1, le=200),
    db: Session = Depends(get_db),
):
    """获取游客打卡列表 — 可按景点筛选"""
    q = db.query(VisitorCheckin).filter(VisitorCheckin.deleted == 0)
    if spot_id:
        q = q.filter(VisitorCheckin.spot_id == spot_id)
    total = q.count()
    rows = q.order_by(desc(VisitorCheckin.created_at)).offset((page - 1) * size).limit(size).all()
    return {
        "items": [
            CheckinOut(
                id=r.id, spot_id=r.spot_id, author=r.author, image=r.image,
                caption=r.caption,
                time=r.created_at.strftime("%Y-%m-%d") if r.created_at else "",
            ).dict() | {
                "created_at": r.created_at.isoformat() if r.created_at else "",
            } for r in rows
        ],
        "total": total, "page": page, "size": size,
    }


@router.post("/checkins")
async def create_checkin(data: CheckinCreate, db: Session = Depends(get_db)):
    """提交新打卡"""
    if not data.caption.strip():
        raise HTTPException(status_code=400, detail="打卡内容不能为空")
    if not data.author.strip():
        data.author = "灵山游客"
    ck = VisitorCheckin(
        spot_id=data.spot_id, author=data.author.strip(), image=data.image,
        caption=data.caption.strip(), created_at=datetime.datetime.now(),
    )
    db.add(ck)
    db.commit()
    db.refresh(ck)
    return {
        "id": ck.id, "spot_id": ck.spot_id, "author": ck.author,
        "image": ck.image, "caption": ck.caption,
        "time": ck.created_at.strftime("%Y-%m-%d") if ck.created_at else "",
    }


@router.delete("/checkins/{checkin_id}")
async def delete_checkin(checkin_id: int, db: Session = Depends(get_db)):
    """删除打卡（管理后台用）"""
    ck = db.query(VisitorCheckin).filter(VisitorCheckin.id == checkin_id, VisitorCheckin.deleted == 0).first()
    if not ck:
        raise HTTPException(status_code=404, detail="打卡不存在")
    ck.deleted = 1
    db.commit()
    return {"deleted": checkin_id, "status": "soft-deleted"}


# ══════════════════════════════════════════════════════════════
# 评价 & 打卡统计（数据大屏用）
# ══════════════════════════════════════════════════════════════

@router.get("/reviews/stats")
async def review_stats(db: Session = Depends(get_db)):
    """评价按景点聚合统计 — 各景点评价数 + 均分 + 总评价/打卡数"""
    from sqlalchemy import func as f

    # 各景点评价统计
    spot_rows = (
        db.query(
            VisitorReview.spot_id,
            f.count(VisitorReview.id).label("cnt"),
            f.round(f.avg(VisitorReview.rating), 1).label("avg_rating"),
        )
        .filter(VisitorReview.deleted == 0)
        .group_by(VisitorReview.spot_id)
        .order_by(desc("cnt"))
        .all()
    )

    spot_stats = [
        {"spot_id": r.spot_id, "count": r.cnt, "avg_rating": float(r.avg_rating) if r.avg_rating else 0}
        for r in spot_rows
    ]

    # 总计
    total_reviews = sum(s["count"] for s in spot_stats)
    total_checkins = db.query(VisitorCheckin).filter(VisitorCheckin.deleted == 0).count()
    avg_all = (
        db.query(f.round(f.avg(VisitorReview.rating), 1))
        .filter(VisitorReview.deleted == 0)
        .scalar()
    )

    return {
        "spots": spot_stats,
        "total_reviews": total_reviews,
        "total_checkins": total_checkins,
        "avg_rating_all": float(avg_all) if avg_all else 0,
    }
