"""
文件上传 → 文档解析 → SQLite 入库 → ChromaDB 向量化 全链路
支持格式: docx / xlsx / txt
"""
import os
import json
import datetime
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Form
from sqlalchemy.orm import Session

from app.db.database import get_db, KnowledgeDoc, get_collection_name
from app.core.rag import add_documents, create_knowledge_base

router = APIRouter()

# 限制
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {".docx", ".xlsx", ".txt"}


# ══════════════════════════════════════════════════════════════
# 文档解析
# ══════════════════════════════════════════════════════════════

def parse_docx(file_bytes: bytes) -> str:
    """从 docx 字节流提取纯文本（含段落和表格）"""
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # 段落
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def parse_xlsx(file_bytes: bytes) -> str:
    """从 xlsx 字节流提取为结构化文本"""
    import io
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"【{sheet_name}】")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
        parts.append("")
    wb.close()
    return "\n".join(parts)


def parse_txt(file_bytes: bytes) -> str:
    """纯文本解析"""
    return file_bytes.decode("utf-8", errors="replace")


def parse_file(filename: str, content: bytes) -> str:
    """根据扩展名选择解析器"""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return parse_docx(content)
    elif ext == ".xlsx":
        return parse_xlsx(content)
    elif ext == ".txt":
        return parse_txt(content)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


# ══════════════════════════════════════════════════════════════
# API
# ══════════════════════════════════════════════════════════════

@router.post("/document")
async def upload_document(
    file: UploadFile = File(...),
    scenic_spot: str = Form(default="灵山胜境"),
    category: str = Form(default="通用"),
    db: Session = Depends(get_db),
):
    """上传文档 → 解析 → SQLite 入库 → ChromaDB 向量化"""

    # ── 校验 ──
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式 '{ext}'，仅支持: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413, detail=f"文件过大 ({len(content)} bytes)，限制 {MAX_FILE_SIZE} bytes"
        )

    # ── 解析 ──
    try:
        text = parse_file(file.filename or "unknown", content)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"文档解析失败: {e}")

    if not text or not text.strip():
        raise HTTPException(status_code=422, detail="文档解析结果为空，请检查文件内容")

    # ── SQLite 入库 ──
    title = os.path.splitext(file.filename or "上传文档")[0]
    doc = KnowledgeDoc(
        title=title,
        content=text,
        category=category,
        scenic_spot=scenic_spot,
        created_at=datetime.datetime.now(),
        updated_at=datetime.datetime.now(),
    )
    db.add(doc)
    db.flush()

    # ── ChromaDB 向量化 ──
    collection = get_collection_name(scenic_spot, db)
    chunk_count = 0
    try:
        chunk_count = add_documents(
            collection,
            [text],
            [{"source": title, "type": "upload", "kb_id": doc.id}],
        )
        coll = create_knowledge_base(collection)
        all_ids = coll.get()["ids"]
        new_ids = all_ids[-chunk_count:] if chunk_count else []
        doc.chroma_ids = json.dumps(new_ids, ensure_ascii=False)
    except Exception:
        chunk_count = 0

    db.commit()
    db.refresh(doc)

    return {
        "id": doc.id,
        "filename": file.filename,
        "title": title,
        "scenic_spot": scenic_spot,
        "category": category,
        "text_length": len(text),
        "chroma_chunks": chunk_count,
        "status": "uploaded_and_indexed",
    }
