"""
Rebuild demo knowledge from the two Word files into SQLite and ChromaDB.

SQLite knowledge_docs is treated as the editable source of truth. This script
upserts entries parsed from the demo package, creates scenic spot slugs, then
rebuilds Chroma collections from SQLite.
"""
import datetime
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document

from app.db.database import KnowledgeDoc, ScenicSpot, SessionLocal
from scripts.sync_knowledge_chroma import sync_collection


DATA_DIR = os.path.abspath(os.path.join(
    os.path.dirname(__file__),
    "..",
    "..",
    "\u793a\u8303\u666f\u533a\u516c\u5f00\u8d44\u6599\u5305",
))
STRUCTURED_FILE = os.path.join(DATA_DIR, "\u7075\u5c71\u80dc\u5883 \u666f\u70b9\u7ed3\u6784\u5316\u6570\u636e\u96c6.docx")
GUIDE_FILE = os.path.join(
    DATA_DIR,
    "\u7075\u5c71\u80dc\u5883\uff1a\u5386\u53f2\u3001\u6587\u5316\u3001\u666f\u70b9\u7279\u8272\u4e0e\u4e2a\u6027\u5316\u6e38\u89c8\u6307\u5357.docx",
)

SCENIC_SLUGS = {
    "\u7075\u5c71\u80dc\u5883": "lingshan",
    "\u62c8\u82b1\u6e7e\u7985\u610f\u5c0f\u9547": "nianhuawan",
}


def ensure_scenic_spots(db):
    for name, slug in SCENIC_SLUGS.items():
        spot = db.query(ScenicSpot).filter(
            (ScenicSpot.name == name) | (ScenicSpot.slug == slug)
        ).first()
        if spot:
            spot.name = name
            spot.slug = slug
            spot.enabled = 1
        else:
            db.add(ScenicSpot(
                name=name,
                slug=slug,
                description=f"{name}\u77e5\u8bc6\u5e93",
                enabled=1,
                created_at=datetime.datetime.now(),
            ))
    db.commit()


def upsert_doc(db, title: str, content: str, category: str, scenic_spot: str):
    title = (title or "").strip()
    content = (content or "").strip()
    if not title or not content:
        return None

    doc = db.query(KnowledgeDoc).filter(
        KnowledgeDoc.title == title,
        KnowledgeDoc.scenic_spot == scenic_spot,
    ).first()
    now = datetime.datetime.now()
    if doc:
        doc.content = content
        doc.category = category
        doc.updated_at = now
    else:
        doc = KnowledgeDoc(
            title=title,
            content=content,
            category=category,
            scenic_spot=scenic_spot,
            created_at=now,
            updated_at=now,
            chroma_ids="[]",
        )
        db.add(doc)
    return doc


def row_values(row):
    return [cell.text.strip() for cell in row.cells]


def import_structured_doc(db):
    doc = Document(STRUCTURED_FILE)
    count = 0
    for table in doc.tables:
        if not table.rows:
            continue
        headers = row_values(table.rows[0])
        if "\u666f\u533a\u540d\u79f0" not in headers or "\u666f\u70b9\u540d\u79f0" not in headers:
            continue
        scenic_idx = headers.index("\u666f\u533a\u540d\u79f0")
        title_idx = headers.index("\u666f\u70b9\u540d\u79f0")
        for row in table.rows[1:]:
            cells = row_values(row)
            if len(cells) <= max(scenic_idx, title_idx):
                continue
            scenic = cells[scenic_idx] or "\u7075\u5c71\u80dc\u5883"
            title = cells[title_idx]
            if not title:
                continue
            parts = [f"\u666f\u70b9\u540d\u79f0\uff1a{title}"]
            for header, value in zip(headers, cells):
                if not value or header in ("\u666f\u533a\u540d\u79f0", "\u666f\u70b9ID", "\u666f\u70b9\u540d\u79f0"):
                    continue
                parts.append(f"\u3010{header}\u3011{value}")
            upsert_doc(db, title, "\n".join(parts), "\u666f\u70b9\u8bb2\u89e3", scenic)
            count += 1
    db.commit()
    return count


def import_guide_paragraphs(db):
    doc = Document(GUIDE_FILE)
    count = 0
    current_title = ""
    current_lines = []

    def flush():
        nonlocal count, current_title, current_lines
        content = "\n".join(current_lines).strip()
        if current_title and content:
            category = "\u8def\u7ebf\u63a8\u8350" if "\u8def\u7ebf" in current_title or "\u8def\u7ebf\u89c4\u5212" in content else "\u6587\u53f2\u8d44\u6599"
            upsert_doc(db, current_title, content, category, "\u7075\u5c71\u80dc\u5883")
            count += 1
        current_title = ""
        current_lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = "Heading" in style_name or (len(text) <= 32 and not text.endswith(("\u3002", "\uff01", "\uff1f")))
        if is_heading:
            flush()
            current_title = text
        else:
            current_lines.append(text)
    flush()
    db.commit()
    return count


def import_guide_tables(db):
    doc = Document(GUIDE_FILE)
    count = 0
    for index, table in enumerate(doc.tables):
        if not table.rows:
            continue
        headers = row_values(table.rows[0])
        rows = [row_values(row) for row in table.rows[1:]]
        if headers[:3] == ["\u7968\u79cd", "\u4ef7\u683c", "\u9002\u7528\u4eba\u7fa4"]:
            title = "\u95e8\u7968\u4e0e\u4f18\u60e0\u653f\u7b56"
            lines = ["\u95e8\u7968\u7968\u4ef7\u4fe1\u606f\uff1a"]
            for cells in rows:
                if len(cells) >= 3:
                    lines.append(f"{cells[0]}\uff1a{cells[1]}\uff0c\u9002\u7528\u4eba\u7fa4\uff1a{cells[2]}")
            upsert_doc(db, title, "\n".join(lines), "\u5b9e\u7528\u4fe1\u606f", "\u7075\u5c71\u80dc\u5883")
            count += 1
        elif headers[:2] == ["\u9879\u76ee", "\u8be6\u7ec6\u4fe1\u606f"]:
            title = f"\u6307\u5357\u8868\u683c-{index + 1}"
            lines = []
            for cells in rows:
                if len(cells) >= 2 and cells[0] and cells[1]:
                    lines.append(f"\u3010{cells[0]}\u3011{cells[1]}")
            if lines:
                upsert_doc(db, title, "\n".join(lines), "\u6587\u53f2\u8d44\u6599", "\u7075\u5c71\u80dc\u5883")
                count += 1
    db.commit()
    return count


def main():
    db = SessionLocal()
    try:
        ensure_scenic_spots(db)
        structured_count = import_structured_doc(db)
        paragraph_count = import_guide_paragraphs(db)
        table_count = import_guide_tables(db)
        chunks = {}
        for scenic_name in SCENIC_SLUGS:
            chunks[scenic_name] = sync_collection(db, scenic_name)
        print(f"Structured rows upserted: {structured_count}")
        print(f"Guide sections upserted: {paragraph_count}")
        print(f"Guide tables upserted: {table_count}")
        print(f"Chroma chunks: {chunks}")
    finally:
        db.close()


if __name__ == "__main__":
    main()
