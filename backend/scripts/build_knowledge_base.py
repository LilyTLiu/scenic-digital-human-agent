"""
知识库构建脚本
从示范景区资料中提取文档内容 → 切片 → 向量化 → 存入ChromaDB
"""
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from app.core.rag import create_knowledge_base, add_documents


def load_docx_text(filepath: str) -> str:
    """从docx文件提取纯文本"""
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def load_docx_tables(filepath: str) -> list[dict]:
    """从docx文件提取表格数据"""
    doc = Document(filepath)
    results = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                results.append(dict(zip(headers, cells)))
    return results


def build_lingshan_kb():
    """构建灵山胜境知识库"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    contest_dir = os.path.abspath(os.path.join(script_dir, "..", ".."))
    data_dir = os.path.join(contest_dir, "示范景区公开资料包")

    # 1. 加载灵山胜境指南
    guide_path = os.path.join(data_dir, "灵山胜境：历史、文化、景点特色与个性化游览指南.docx")
    guide_text = load_docx_text(guide_path)

    # 2. 加载结构化景点数据
    dataset_path = os.path.join(data_dir, "灵山胜境 景点结构化数据集.docx")
    spots = load_docx_tables(dataset_path)

    # 构建每条景点的完整文本
    spot_texts = []
    for spot in spots:
        fields = []
        for key, val in spot.items():
            if val and key not in ("备注",):
                fields.append(f"{key}：{val}")
        spot_texts.append("\n".join(fields))

    # 3. 导入知识库
    collection_name = "lingshan"
    print(f"Creating collection: {collection_name}")

    # 导入指南文档
    guide_chunks = add_documents(
        collection_name,
        [guide_text],
        [{"source": "灵山胜境指南", "type": "guide"}],
    )
    print(f"  Guide: {guide_chunks} chunks added")

    # 导入景点结构化数据
    spot_chunks = add_documents(
        collection_name,
        spot_texts,
        [{"source": "景点结构化数据集", "type": "spot", "spot_name": s.get("景点名称", "")} for s in spots],
    )
    print(f"  Spots: {spot_chunks} chunks added")

    print(f"\nTotal documents in collection: {guide_chunks + spot_chunks}")
    print("Knowledge base built successfully!")


if __name__ == "__main__":
    build_lingshan_kb()
